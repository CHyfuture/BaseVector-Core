"""
Word文档解析器
"""
from pathlib import Path
from typing import Any, Dict

try:
    from docx import Document
    from docx.document import Document as DocumentType
except ImportError:
    Document = None

from ability.operators.parsers.base_parser import BaseParser


class DocxParser(BaseParser):
    """Word文档解析器，使用python-docx"""

    def __init__(self, config: Dict[str, Any] = None):
        """
        初始化Word解析器

        Args:
            config: 配置字典
        """
        super().__init__(config)
        self.supported_extensions = [".docx", ".doc"]

        if Document is None:
            raise ImportError(
                "python-docx is required for Word parsing. "
                "Install it with: pip install python-docx"
            )

    def _initialize(self) -> None:
        """初始化Word解析器"""
        if Document is None:
            raise ImportError("python-docx is not installed")
        self.logger.info("Word parser initialized")

    def _parse(self, file_path: Path, **kwargs) -> Dict[str, Any]:
        """
        解析Word文档

        Args:
            file_path: Word文件路径
            **kwargs: 额外的处理参数

        Returns:
            解析结果字典
        """
        doc = Document(str(file_path))

        # 提取元数据
        core_props = doc.core_properties
        metadata = {
            "title": core_props.title or "",
            "author": core_props.author or "",
            "subject": core_props.subject or "",
            "keywords": core_props.keywords or "",
            "comments": core_props.comments or "",
            "created": str(core_props.created) if core_props.created else "",
            "modified": str(core_props.modified) if core_props.modified else "",
        }

        content_parts = []
        structure = {
            "paragraphs": [],
            "tables": [],
        }

        # 提取段落
        for para_idx, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            if not text:
                continue

            # 判断段落样式
            style = paragraph.style.name if paragraph.style else "Normal"
            is_heading = style.startswith("Heading")

            if is_heading:
                # 提取标题级别
                level = 1
                if "Heading" in style:
                    try:
                        level = int(style.replace("Heading", "").strip())
                    except ValueError:
                        level = 1

                content_parts.append(f"{'#' * level} {text}\n")
                structure["paragraphs"].append(
                    {
                        "index": para_idx,
                        "type": "heading",
                        "level": level,
                        "text": text,
                    }
                )
            else:
                content_parts.append(f"{text}\n\n")
                structure["paragraphs"].append(
                    {
                        "index": para_idx,
                        "type": "paragraph",
                        "text": text,
                    }
                )

        # 提取表格
        for table_idx, table in enumerate(doc.tables):
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)

            if table_data:
                markdown_table = self._table_to_markdown(table_data)
                content_parts.append(f"\n### Table {table_idx + 1}\n\n{markdown_table}\n\n")
                structure["tables"].append(
                    {
                        "table_index": table_idx,
                        "rows": len(table_data),
                        "cols": len(table_data[0]) if table_data else 0,
                    }
                )

        content = "\n".join(content_parts)

        return {
            "content": content,
            "metadata": metadata,
            "structure": structure,
        }

    def _table_to_markdown(self, table_data: list) -> str:
        """
        将表格数据转换为Markdown格式

        Args:
            table_data: 表格数据（二维列表）

        Returns:
            Markdown表格字符串
        """
        if not table_data:
            return ""

        markdown_lines = []

        # 表头
        if table_data:
            header = table_data[0]
            markdown_lines.append("| " + " | ".join(str(cell) for cell in header) + " |")
            markdown_lines.append("| " + " | ".join("---" for _ in header) + " |")

        # 表体
        for row in table_data[1:]:
            markdown_lines.append("| " + " | ".join(str(cell) for cell in row) + " |")

        return "\n".join(markdown_lines)
